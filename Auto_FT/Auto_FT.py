hub="C:/Users/xingxingyu/Desktop/FT_hub"	# 修改路径分隔符为 /

#使用input接收一串字符串，存放到LiaoHao变量中
LiaoHao=input("请输入料号:")
#分割料号，取出XPD后的三位数字，如701，702，903
#如果option=701，则打开hub路径下“701”文件，如果option=702，则打开hub路径下“702”文件，如果option=903，则打开hub路径下“903”文件
option=int(LiaoHao[3:6])


import pyperclip
if option==701:
    filename_target="YX50081A_ XPD701DPS6545VUY_FT测试规范_V1_20250807.docx"
    input("请先将701_Reg数据复制到粘贴板,然后点击回车。")
    clip=pyperclip.paste()
    #把clip中的内容按顺序划分为48个两位数字，然后存入一个含有48个元素的数组array
    array=clip.split()
    #打印”701_Reg数据已接受，请点击回车。（换行）内容如下：“
    print("701_Reg数据已接受,请点击回车。\n内容如下:\n",clip)
elif option==702:
    filename_target="YX50081A_ XPD702DP6520BVUY2_FT测试规范_V1_20250807.docx"
    input("请先将702_Reg数据复制到粘贴板,然后点击回车。")
    clip=pyperclip.paste()
    #把clip中的内容按顺序划分为48个两位数字，然后存入一个含有48个元素的数组array
    array=clip.split()
    #打印”702_Reg数据已接受，请点击回车。（换行）内容如下：“
    print("702_Reg数据已接受,请点击回车。\n内容如下:\n",clip)
elif option==903:
    filename_target="YX50081A_XPD903DP45A4_FT测试规范_V1_20250820.docx"
    input("请先将C1口数据复制到粘贴板,然后点击回车。")
    clip=pyperclip.paste()
    #把clip中的内容按顺序划分为48个两位数字，然后存入一个含有48个元素的数组array
    array=clip.split()
    #打印”C1口数据已接受，请复制C2口数据后点击回车。（换行）C1口内容如下：“
    print("C1口数据已接受,请复制C2口数据后点击回车。\nC1口内容如下:\n",clip)
    #暂停程序，等待用户输入
    input()
    clip2=pyperclip.paste()
    array2=clip2.split()
    print("C2口数据已接受,请复制C2口数据后点击回车。\nC2口内容如下:\n",clip2)
    
import os
import docx
import datetime
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
temp=datetime.datetime.now()
data=temp.strftime("%Y%m%d")
filename_save="YX50081A_"+LiaoHao+"_FT测试规范_V1_"+data+".docx"

#打开hub路径下的目标文档
doc=docx.Document(os.path.join(hub,filename_target))
#获取第一个表格
table=doc.tables[0]

'''
#提取并打印表头
headers=[cell.text for cell in table.rows[0].cells]
print("表头:",headers)
'''
#提取表格数据，将第一行的第四个数据改为_修订，然后写回文件，最后打印第一行数据
for row in table.rows[1:2]:
    row_data=[cell.text for cell in row.cells]
    row_data[3]=data
    print("第一行数据改为:",row_data)
    #exit()      #调试点1，调试用，注释掉后正常运行

    #将第一行数据写回文件
    for i in range(0,4):
        row.cells[i].text=row_data[i]



#打印N~N行的内容，格式为：
'''
行号 【内容】
1 第一行内容
2 第二行内容
'''

if option==701:
    '''    
    #逐行打印1~196行的内容,用于调试
    for line_num in range(1, 197):
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
    '''
    #打印"701寄存器数据如下："
    print("701寄存器数据如下:")
    for line_num in range(141, 189):  # 141-189共48行
        # 查找第一个等号的位置
        first_equal = doc.paragraphs[line_num].text.find('=')
        # 从第一个等号之后的位置开始查找第二个等号
        pos = doc.paragraphs[line_num].text.find('=', first_equal + 1)
        # 计算array索引（line_num-141对应array[0]到array[47]）
        array_index = line_num - 141
        # 将第二个等号后两位数字改为array[array_index]
        doc.paragraphs[line_num].text = doc.paragraphs[line_num].text[:pos+1] + array[array_index] + doc.paragraphs[line_num].text[pos+3:]
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
elif option==702:
#162-209是702的寄存器位置
    '''
    #逐行打印1~196行的内容,用于调试
    for line_num in range(1, 197):
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
    '''
    #打印"702寄存器数据如下："
    print("702寄存器数据如下:")
    for line_num in range(162, 210):  # 162-210共48行
        # 查找第一个等号的位置
        first_equal = doc.paragraphs[line_num].text.find('=')
        # 从第一个等号之后的位置开始查找第二个等号
        pos = doc.paragraphs[line_num].text.find('=', first_equal + 1)
        # 计算array索引（line_num-162对应array[0]到array[47]）
        array_index = line_num - 162
        # 将第二个等号后两位数字改为array[array_index]
        doc.paragraphs[line_num].text = doc.paragraphs[line_num].text[:pos+1] + array[array_index] + doc.paragraphs[line_num].text[pos+3:]
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
elif option==903:
#146-193是C1,197-244是C2
# 遍历146-193行，每行找到第二个等号位置，将后两位数字替换为array中对应元素

    #打印"C1口数据如下："
    print("C1口数据如下:")
    for line_num in range(146, 194):  # 146-193共48行
        # 查找第一个等号的位置
        first_equal = doc.paragraphs[line_num].text.find('=')
        # 从第一个等号之后的位置开始查找第二个等号
        pos = doc.paragraphs[line_num].text.find('=', first_equal + 1)
        # 计算array索引（line_num-146对应array[0]到array[47]）
        array_index = line_num - 146
        # 将第二个等号后两位数字改为array[array_index]
        doc.paragraphs[line_num].text = doc.paragraphs[line_num].text[:pos+1] + array[array_index] + doc.paragraphs[line_num].text[pos+3:]
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
    #打印"C2口数据如下："
    print("C2口数据如下:")
    #与146-194行同样的逻辑，使用array2替换对应
    for line_num in range(197, 245):  # 197-244共48行
        first_equal = doc.paragraphs[line_num].text.find('=')
        pos = doc.paragraphs[line_num].text.find('=', first_equal + 1)
        array_index = line_num - 197
        doc.paragraphs[line_num].text = doc.paragraphs[line_num].text[:pos+1] + array2[array_index] + doc.paragraphs[line_num].text[pos+3:]
        print(line_num,"【",doc.paragraphs[line_num].text,"】")
    #用run方法分别把146-193行和197-244行字体设置为12pt  
    for line_num in range(146, 194):  # 146-193共48行
        for run in doc.paragraphs[line_num].runs:
            run.font.size=docx.shared.Pt(12)
    for line_num in range(197, 245):  # 197-244共48行
        for run in doc.paragraphs[line_num].runs:
            run.font.size=docx.shared.Pt(12)

#exit()              #调试点2,调试用，注释掉后正常运行


#保存文档
doc.save(os.path.join(hub,filename_save))
print("完成")
input("按回车键退出")
exit()