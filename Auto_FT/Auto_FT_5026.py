
hub_save=r"C:\Users\chenqinyuan\Desktop\8"	# 修改路径分隔符为 /
#使用input接收一串字符串，存放到LiaoHao变量中
LiaoHao=input("请输入料号:")
filename_target=""
#分割料号，取出XPD后的三位数字，如721
#如果option=721，则打开hub路径下“721”文件，如果option=721，则打开hub路径下“721”文件，如果option=721，则打开hub路径下“721”文件
option=int(LiaoHao[3:6])#取出料号的第4位到第6位，如721
#根据option打开对应的模板文档，只能用if语句
available_options=[719,720,721,722,723,729,730,736,737,738,740,741,750,767,902]

if option not in available_options:
    raise ValueError("未定义的系列！")
elif option==719:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD719系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD719BP25Q_FT测试规范_V1_20250703.docx"
    input("请先将719_Reg数据复制到粘贴板,然后点击回车。")
elif option==720:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD720系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD720A_FT测试模板_V1_20230814.docx"
    input("请先将720_Reg数据复制到粘贴板,然后点击回车。")
elif option==721:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD721系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD721A_FT测试模板_V1_20230814.docx"
    input("请先将721_Reg数据复制到粘贴板,然后点击回车。")
elif option==722:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD722系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD722D65_FT测试规范_V1_20250724.docx"
    input("请先将722_Reg数据复制到粘贴板,然后点击回车。")
elif option==723:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD723系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026B_XPD723A27K_FT测试规范_V1_20231229.docx"
    input("请先将723_Reg数据复制到粘贴板,然后点击回车。")
elif option==729:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD729系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD729APS30_FT测试规范_V1_20250305.docx"
    input("请先将729_Reg数据复制到粘贴板,然后点击回车。")
elif option==730:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD730系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD730DP45DCB_FT测试规范_V1_20250728.docx"
    input("请先将730_Reg数据复制到粘贴板,然后点击回车。")   
elif option==736:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD736系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD736APS25B_FT测试规范_V1_20240830.docx"
    input("请先将736_Reg数据复制到粘贴板,然后点击回车。")       
elif option==737:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD737系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD737C30R_FT测试规范_V1_20250728.docx"
    input("请先将737_Reg数据复制到粘贴板,然后点击回车。")       
elif option==738:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD738系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD738BPUK_FT测试规范_V1_20250521.docx"
    input("请先将738_Reg数据复制到粘贴板,然后点击回车。")       
elif option==740:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD740系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD740APS25_FT测试规范_V1_20231208.docx"
    input("请先将740_Reg数据复制到粘贴板,然后点击回车。")       
elif option==741:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD741系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD741BPS25V_FT测试规范_V1_20250623.docx"
    input("请先将741_Reg数据复制到粘贴板,然后点击回车。")       
elif option==750:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD750系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD750A18_FT测试规范_V1_20230520.docx"
    input("请先将750_Reg数据复制到粘贴板,然后点击回车。")       
elif option==767:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD767系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD767A_FT测试规范_V1_20240516.docx"
    input("请先将767_Reg数据复制到粘贴板,然后点击回车。")       
elif option==768:
    hub_open=r""	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD768A_FT测试模板_V1_20230814.docx"
    input("请先将768_Reg数据复制到粘贴板,然后点击回车。")       
elif option==902:
    hub_open=r"\\192.168.212.104\h\project\usb controller\YX5026\10_量产测试\FT测试规范\YX5026A_XPD902C系列FT测试规范"	# 修改路径分隔符为 /
    filename_target="YX5026A_XPD902CBP_FT测试规范_V1_20240605.docx"
    input("请先将902_Reg数据复制到粘贴板,然后点击回车。")       




import os
import docx
import datetime
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
temp=datetime.datetime.now()
data=temp.strftime("%Y%m%d")
filename_save="YX5026A_"+LiaoHao+"_FT测试规范_V1_"+data+".docx"
print("filename_target:",filename_target)

#打开hub路径下的目标文档
doc=docx.Document(os.path.join(hub_open,filename_target))
#获取第一个表格
table=doc.tables[0]

'''
#提取并打印表头
headers=[cell.text for cell in table.rows[0].cells]
print("表头:",headers)

'''
#提取表格数据，将第一行的第四个数据改为修订日期，然后写回文件，最后打印第一行数据
for row in table.rows[1:2]:
    row_data=[cell.text for cell in row.cells]
    row_data[3]=data
    row_data[4]="陈钦源"
    print("修订表第一行数据改为:",row_data)
    #exit()      #调试点1，调试用，注释掉后正常运行

    #将第一行数据写回文件
    for i in range(0,5):
        row.cells[i].text=row_data[i]
        
#获取文档总行数
num_line=len(doc.paragraphs)
#使用一个循环，将每一行以字符串的形式存于一个数组中，并将待修改的目标内容存在第二个数组中
line_array=[]
line_array2=[]
for i in range(0,num_line):
    line_array.append(doc.paragraphs[i].text)
#遍历数组line_array，如果元素以"0x"开头，且元素中包含"即："，则用变量pos1记录"即："位置，并从"即："后继续查找是否存在等号
for i in range(0,len(line_array)):
    if line_array[i].startswith("0x"):
        if line_array[i].find("即：")>-1:
            pos1=line_array[i].find("即：")
            if line_array[i].find("=",pos1+1)>-1:
                line_array2.append(line_array[i])
                pos2=line_array[i].find("=",pos1+1) 
                line_array2.append(line_array[i])
#找到目标内容的开始行号
for i in range(0,len(line_array)):
    if line_array[i].startswith("0x"):
        if line_array[i].find("即：")>-1:
            pos1=line_array[i].find("即：")
            if line_array[i].find("=",pos1+1)>-1:
                begin_line=i
                break

'''
#打印dec的第begin_line行
print(doc.paragraphs[begin_line].text)
exit()
'''

'''
#打印line_array2，格式为：数组下标  内容  等号位置  第二个等号位置
for i in range(0,len(line_array2)):
    print(i," ",line_array2[i]," ",line_array2[i].find("=")," ",line_array2[i].find("=",line_array2[i].find("=")+1))

'''
import pyperclip
clip=pyperclip.paste()
#把clip中的内容按顺序划分为32个两位数字，然后存入一个含有32个元素的数组array
array=clip.split()
#打印从剪切板中获取的字符串
print("选定的寄存器数据:\n", clip,"\n")
print("插入到文档效果如下:\n")


# 用一个32次循环,在循环内，查找从begin_line开始的32行的每一行的第二个"="的位置，并删除该位置后的两个字符，再插入array[i],不用run
for i in range(0,32):
    #在doc.paragraphs[begin_line+i].text内定位第二个"="的位置
    pos=doc.paragraphs[begin_line+i].text.find("=",doc.paragraphs[begin_line+i].text.find("=")+1)
    #使用run方法删除第二个"="后的两个字符，在此处插入array[i]，且保持文本格式不变.
    doc.paragraphs[begin_line+i].text=doc.paragraphs[begin_line+i].text[:pos+1]+array[i]+doc.paragraphs[begin_line+i].text[pos+3:]
    #用run方法设置这一行字体大小为12pt
    for run in doc.paragraphs[begin_line+i].runs:
        run.font.size=docx.shared.Pt(12)
    #打印修改后的内容，格式为：行号 【内容】
    print("行号:",begin_line+i," ",doc.paragraphs[begin_line+i].text)


'''
#从begin_line开始打印32行，格式为：行号 【内容】
for i in range(0,32):
    print("行号:",begin_line+i," ",doc.paragraphs[begin_line+i].text)

'''

#exit()              #调试点2,调试用，注释掉后正常运行


#保存文档
doc.save(os.path.join(hub_save,filename_save))
print("完成，已保存为：",filename_save)
exit()